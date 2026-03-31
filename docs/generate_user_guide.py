"""Generate CAGD Website User Guide (.docx)"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page numbering helper ───────────────────────────────────────────
def add_page_numbers(doc):
    """Add page numbers to footer, starting from second page (cover page excluded)."""
    # Set the document to start page numbering from 0 so first content page = 1
    for section in doc.sections:
        section.different_first_page_header_footer = True
        # Footer for pages after the first
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.clear()

        # "Page " text
        run1 = p.add_run('Page ')
        run1.font.size = Pt(9)
        run1.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # PAGE field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run2 = p.add_run()
        run2._r.append(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run3 = p.add_run()
        run3._r.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run4 = p.add_run()
        run4._r.append(fldChar2)

        # Header with document title
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.clear()
        hr = hp.add_run('CAGD Website — User Guide')
        hr.font.size = Pt(8)
        hr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        hr.italic = True

        # First page header/footer are empty (cover page)
        first_footer = section.first_page_footer
        first_header = section.first_page_header
        if first_footer.paragraphs:
            first_footer.paragraphs[0].clear()
        if first_header.paragraphs:
            first_header.paragraphs[0].clear()

# ── Styles ──────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x00, 0x52, 0x33)  # Ghana green
    if level == 1:
        hs.font.size = Pt(22)
    elif level == 2:
        hs.font.size = Pt(16)
    else:
        hs.font.size = Pt(13)

# ── Title Page ──────────────────────────────────────────────────────
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('CONTROLLER & ACCOUNTANT-GENERAL\'S DEPARTMENT')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x00, 0x52, 0x33)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Official Website — User Guide')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)  # Gold

doc.add_paragraph()

url = doc.add_paragraph()
url.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = url.add_run('https://cagd.gov.gh')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x00, 0x52, 0x33)

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('March 2026\nVersion 1.0\nICT Division')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ── Table of Contents placeholder ───────────────────────────────────
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Introduction',
    '2. Getting Started',
    '3. Homepage',
    '4. About CAGD',
    '5. Management & Leadership',
    '6. Divisions',
    '7. Projects & Reforms',
    '8. News & Media',
    '9. Events',
    '10. Reports & Publications',
    '11. Gallery',
    '12. Staff Directory',
    '13. Forms Library',
    '14. Regional Offices',
    '15. System Status',
    '16. Contact & Support',
    '17. FAQ',
    '18. ChatBot Assistant',
    '19. Accessibility Features',
    '20. Keyboard Shortcuts',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.runs[0].font.size = Pt(11)

doc.add_page_break()

# ── Helper ──────────────────────────────────────────────────────────
def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run(f'💡 {text}')
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(f' — {text}')
    else:
        p.add_run(text)

# ── 1. Introduction ─────────────────────────────────────────────────
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    'The Controller & Accountant-General\'s Department (CAGD) website at cagd.gov.gh serves as the official digital '
    'platform for the department. It provides public access to news, events, reports, forms, staff information, and '
    'details about CAGD\'s operations across all 16 regions of Ghana.'
)
doc.add_paragraph(
    'This guide covers all features available to the public and staff when using the website.'
)

doc.add_heading('Key Features at a Glance', level=2)
features = [
    ('Latest News & Press Releases', 'Stay updated with CAGD announcements and publications'),
    ('Event Registration', 'Register for CAGD events with QR code confirmation'),
    ('Reports & Publications', 'Download annual reports, financial statements, and publications'),
    ('Forms Library', 'Access and download official CAGD forms (payroll, pension, GIFMIS, etc.)'),
    ('Staff Directory', 'Search and find CAGD staff by name, division, or department'),
    ('Regional Offices Map', 'Interactive map showing all 16 regional offices with contact details'),
    ('System Status', 'Check real-time status of GIFMIS, e-Payslip, TPRS, and CAGD Website'),
    ('Photo Gallery', 'Browse albums of CAGD events and activities'),
    ('ChatBot Assistant', 'Get quick answers to common questions'),
    ('Dark Mode & Language', 'Switch between light/dark themes and English/Twi languages'),
]
for title, desc in features:
    add_bullet(desc, title)

# ── 2. Getting Started ──────────────────────────────────────────────
doc.add_heading('2. Getting Started', level=1)
doc.add_heading('Accessing the Website', level=2)
doc.add_paragraph('Open your web browser (Chrome, Firefox, Edge, Safari) and navigate to:')
p = doc.add_paragraph()
run = p.add_run('https://cagd.gov.gh')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x00, 0x52, 0x33)

doc.add_paragraph('The website is fully responsive and works on desktops, tablets, and mobile phones.')

doc.add_heading('Browser Requirements', level=2)
add_bullet('Google Chrome 90+')
add_bullet('Mozilla Firefox 90+')
add_bullet('Microsoft Edge 90+')
add_bullet('Safari 14+')
add_bullet('Stable internet connection')

doc.add_heading('Navigation', level=2)
doc.add_paragraph(
    'The main navigation menu is located at the top of every page. On desktop, hover over menu items to see '
    'dropdown sub-menus. On mobile devices, tap the hamburger menu icon (☰) in the top-right corner to '
    'open the navigation panel.'
)
doc.add_paragraph('The navigation menu includes:')
nav_items = [
    ('Home', 'Returns to the homepage'),
    ('About', 'Who We Are, Mission & Vision, Structure, History, Core Functions'),
    ('Services', 'Core Functions, Finance Administration, Treasury, FMS, Payroll, ICT, Audit'),
    ('Resources', 'Reports, Gallery, FAQ, Projects, News'),
    ('News & Events', 'Latest News, Press Releases, CAGD Digest, Treasury News, Events'),
    ('Contact', 'Contact form and office information'),
]
for title, desc in nav_items:
    add_bullet(desc, title)

# ── 3. Homepage ─────────────────────────────────────────────────────
doc.add_heading('3. Homepage', level=1)
doc.add_paragraph('The homepage is the main landing page and provides quick access to all key information.')

doc.add_heading('Hero Slider', level=2)
doc.add_paragraph(
    'The top of the homepage features a large animated slideshow highlighting CAGD\'s mission, current events, '
    'and key announcements. Each slide may include a title, subtitle, and a call-to-action button that links to '
    'relevant pages.'
)

doc.add_heading('Announcement Banner', level=2)
doc.add_paragraph(
    'When active, a coloured banner appears at the very top of the page with important notices '
    '(e.g., system maintenance schedules, deadline reminders). You can dismiss it by clicking the X button. '
    'The banner will reappear on your next visit.'
)

doc.add_heading('Statistics Section', level=2)
doc.add_paragraph(
    'Animated counters display key figures such as the number of news articles, reports published, '
    'events held, and gallery albums. These numbers update automatically as content is added.'
)

doc.add_heading('Core Functions', level=2)
doc.add_paragraph(
    'A visual card section highlighting CAGD\'s eight core functions: Payroll Management, Pension Administration, '
    'Revenue Management, Expenditure Control, Government Accounting, Financial Reporting, Systems Management, '
    'and Banking Services.'
)

doc.add_heading('e-Services Quick Links', level=2)
doc.add_paragraph('Direct links to key government e-services:')
add_bullet('GIFMIS — Ghana Integrated Financial Management Information System')
add_bullet('e-Payslip — Government electronic payslip portal')
add_bullet('TPRS — Treasury Payment Receipt System')
add_bullet('GoGePV — Government of Ghana e-Procurement Voucher')

doc.add_heading('Latest News & Events', level=2)
doc.add_paragraph(
    'The homepage displays the most recent news articles and upcoming events. Click on any item to read the '
    'full article or view event details.'
)

doc.add_heading('Newsletter Subscription', level=2)
doc.add_paragraph(
    'At the bottom of the homepage (and in the footer), you can enter your email address to subscribe to '
    'CAGD\'s newsletter and receive updates about news, events, and announcements.'
)

doc.add_heading('Partners', level=2)
doc.add_paragraph('A scrolling carousel showcasing CAGD\'s partner organisations and their logos.')

# ── 4. About CAGD ──────────────────────────────────────────────────
doc.add_heading('4. About CAGD', level=1)

doc.add_heading('Who We Are', level=2)
doc.add_paragraph(
    'Located at About → Who We Are, this page provides a comprehensive overview of the Controller & '
    'Accountant-General\'s Department, including its mandate, role within the Ministry of Finance, and '
    'organisational values.'
)

doc.add_heading('Mission & Vision', level=2)
doc.add_paragraph('Displays the department\'s mission statement, vision, and core values.')

doc.add_heading('Organisational Structure', level=2)
doc.add_paragraph(
    'A visual representation of CAGD\'s organisational hierarchy, showing the relationship between '
    'divisions, departments, and regional offices.'
)

doc.add_heading('Our History', level=2)
doc.add_paragraph('A timeline of CAGD\'s establishment and key milestones in its history.')

doc.add_heading('Core Functions', level=2)
doc.add_paragraph(
    'Detailed descriptions of CAGD\'s eight core functions with explanations of each function\'s '
    'responsibilities and scope.'
)

# ── 5. Management & Leadership ──────────────────────────────────────
doc.add_heading('5. Management & Leadership', level=1)

doc.add_heading('Leadership Profiles', level=2)
doc.add_paragraph(
    'Located at Management → Leadership, this page showcases the Controller & Accountant-General, '
    'Deputy Controllers, and other senior management staff. Each profile includes a photo, name, '
    'title, and biography. Click on any profile to view the full details.'
)

doc.add_heading('Regional Directors', level=2)
doc.add_paragraph(
    'Located at Management → Regional Directors, this page lists the directors of all 16 CAGD '
    'regional offices. Each entry includes the director\'s name, photo, region, and contact details.'
)

# ── 6. Divisions ────────────────────────────────────────────────────
doc.add_heading('6. Divisions', level=1)
doc.add_paragraph('CAGD operates through six main divisions, each with a dedicated page on the website:')

divisions = [
    ('Finance & Administration', 'Manages CAGD\'s internal finance, human resources, procurement, and administrative operations.'),
    ('Treasury Management', 'Responsible for government banking, cash management, and the Consolidated Fund.'),
    ('Financial Management Systems (FMS)', 'Oversees GIFMIS implementation, system development, and financial technology.'),
    ('Information & Communication Technology (ICT)', 'Manages IT infrastructure, digital services, cybersecurity, and technical support.'),
    ('Payroll', 'Handles government payroll processing, salary validation, and payroll audits for all public sector employees.'),
    ('Audit', 'Conducts internal audits, compliance reviews, and financial inspections across government agencies.'),
]
for title, desc in divisions:
    add_bullet(desc, title)

# ── 7. Projects ─────────────────────────────────────────────────────
doc.add_heading('7. Projects & Reforms', level=1)
doc.add_paragraph('CAGD leads several major reform initiatives:')

doc.add_heading('PFMRP — Public Financial Management Reform Programme', level=2)
doc.add_paragraph(
    'Ghana\'s comprehensive public financial management reform programme aimed at improving fiscal discipline, '
    'budget credibility, and accountability in the use of public resources.'
)

doc.add_heading('IPSAS — International Public Sector Accounting Standards', level=2)
doc.add_paragraph(
    'Ghana\'s transition to international accounting standards for the public sector, ensuring transparency '
    'and comparability of government financial reports with international best practices.'
)

# ── 8. News & Media ─────────────────────────────────────────────────
doc.add_heading('8. News & Media', level=1)

doc.add_heading('Latest News', level=2)
doc.add_paragraph(
    'Located at News & Events → All News, this page shows all CAGD news articles. You can filter by category '
    '(General, Announcements, IPSAS, Payroll, GIFMIS, Training, Reforms, Press Release, Events, Digest, Treasury News) '
    'and search by keyword.'
)

doc.add_heading('Press Releases', level=2)
doc.add_paragraph('Official CAGD press releases and media statements.')

doc.add_heading('CAGD Digest', level=2)
doc.add_paragraph(
    'The CAGD Digest is a periodic publication summarising department activities, achievements, and upcoming plans. '
    'Each issue is displayed with its cover and can be read online or downloaded.'
)

doc.add_heading('Treasury News', level=2)
doc.add_paragraph('News and updates specifically related to Treasury Management operations.')

doc.add_heading('Reading a News Article', level=2)
doc.add_paragraph('Click on any news item to view the full article. Each article page includes:')
add_bullet('Full article text with images')
add_bullet('Publication date and author')
add_bullet('Category tags')
add_bullet('Social media sharing buttons (Facebook, Twitter, LinkedIn)')
add_bullet('Copy link button')
add_bullet('Downloadable attachments (if available)')
add_bullet('Related news articles')
add_bullet('Breadcrumb navigation to return to the news list')

doc.add_heading('Newsletter Subscription', level=2)
doc.add_paragraph(
    'On the news page, you can subscribe to receive email notifications when new articles are published. '
    'Enter your name and email address in the subscription form.'
)

# ── 9. Events ───────────────────────────────────────────────────────
doc.add_heading('9. Events', level=1)

doc.add_heading('Browsing Events', level=2)
doc.add_paragraph(
    'Located at News & Events → Events, this page lists all CAGD events. You can view events in a list format '
    'or switch to a calendar view using the toggle button. Use the search bar to find specific events.'
)

doc.add_heading('Event Details', level=2)
doc.add_paragraph('Click on any event to view its full details, including:')
add_bullet('Event title and description')
add_bullet('Date and time (start and end dates)')
add_bullet('Venue/location')
add_bullet('Countdown timer showing days, hours, minutes, and seconds until the event')
add_bullet('Total number of registered participants')
add_bullet('Share buttons')

doc.add_heading('Event Registration', level=2)
doc.add_paragraph(
    'For events that accept registrations, you will see a registration form on the event page. '
    'Fill in the following details:'
)
add_bullet('First Name, Middle Name (optional), and Surname')
add_bullet('Gender')
add_bullet('Participant Type (e.g., government official, private sector, student)')
add_bullet('Region (select from all 16 regions of Ghana)')
add_bullet('Department')
add_bullet('Phone number and Email address')
add_bullet('Organisation')

doc.add_paragraph(
    'After successful registration, you will receive:'
)
add_bullet('A QR code that serves as your registration confirmation — you can save or print this')
add_bullet('A confirmation email with your registration details (if email service is configured)')

doc.add_heading('Event Feedback', level=2)
doc.add_paragraph(
    'After attending an event, you can submit feedback by rating the event (1-5 stars) and leaving '
    'a written comment. This helps CAGD improve future events.'
)

# ── 10. Reports ─────────────────────────────────────────────────────
doc.add_heading('10. Reports & Publications', level=1)
doc.add_paragraph(
    'Located at Resources → Reports, this page provides access to CAGD\'s official reports and publications. '
    'Reports can be filtered by category and searched by keyword.'
)
doc.add_paragraph('Each report entry shows:')
add_bullet('Report title and description')
add_bullet('Category')
add_bullet('Publication date')
add_bullet('File size')
add_bullet('Download count')
doc.add_paragraph('Click the download button to save the report to your device.')

# ── 11. Gallery ─────────────────────────────────────────────────────
doc.add_heading('11. Gallery', level=1)
doc.add_paragraph(
    'The photo gallery showcases images from CAGD events, activities, and official functions. '
    'Photos are organised into albums.'
)
doc.add_paragraph('Features:')
add_bullet('Browse albums with cover photos and photo counts')
add_bullet('Click an album to view all photos')
add_bullet('Click any photo to open a full-screen lightbox viewer')
add_bullet('Navigate between photos using arrow keys or on-screen buttons')
add_bullet('Albums may contain sub-albums for better organisation')

# ── 12. Staff Directory ─────────────────────────────────────────────
doc.add_heading('12. Staff Directory', level=1)
doc.add_paragraph(
    'Located at the Staff Directory page, this feature allows you to search for CAGD staff members. '
    'You can search by name, job title, division, or department.'
)
doc.add_paragraph('Each staff entry displays:')
add_bullet('Full name and photo')
add_bullet('Job title')
add_bullet('Division and department')
add_bullet('Phone number')
add_bullet('Email address')
doc.add_paragraph('Use the search bar at the top to quickly find a specific staff member.')

# ── 13. Forms Library ───────────────────────────────────────────────
doc.add_heading('13. Forms Library', level=1)
doc.add_paragraph(
    'Located at Resources → Forms, this page provides access to official CAGD forms that can be '
    'downloaded, printed, and filled out.'
)
doc.add_paragraph('Forms are organised into categories:')
add_bullet('Payroll forms')
add_bullet('Pension forms')
add_bullet('GIFMIS forms')
add_bullet('HR forms')
add_bullet('Finance forms')
add_bullet('General forms')

doc.add_paragraph(
    'Each form shows its name, form code (e.g., CAGD-001), description, file size, and download count. '
    'Use the search bar or category tabs to find the form you need, then click Download.'
)

# ── 14. Regional Offices ────────────────────────────────────────────
doc.add_heading('14. Regional Offices', level=1)
doc.add_paragraph(
    'Located at Contact → Regional Offices, this page displays an interactive map of all 16 CAGD '
    'regional offices across Ghana.'
)
doc.add_paragraph('Features:')
add_bullet('Interactive map with clickable markers for each regional office')
add_bullet('Click a marker to see the office\'s address, phone number, email, and regional director')
add_bullet('Below the map, a card grid lists all offices with full contact details')
add_bullet('Regional director name and photo for each office')

# ── 15. System Status ───────────────────────────────────────────────
doc.add_heading('15. System Status', level=1)
doc.add_paragraph(
    'Located at the System Status page (/status), this page shows the real-time operational status '
    'of key government systems:'
)
add_bullet('GIFMIS — Ghana Integrated Financial Management Information System')
add_bullet('e-Payslip Portal — Electronic payslip service')
add_bullet('TPRS — Treasury Payment Receipt System')
add_bullet('CAGD Website — This website')

doc.add_paragraph('Status indicators:')
add_bullet('Green — Operational (all systems normal)', 'Operational')
add_bullet('Amber — Degraded Performance (system is slow or partially unavailable)', 'Degraded')
add_bullet('Red — Service Outage (system is down)', 'Outage')
add_bullet('Blue — Under Maintenance (planned maintenance in progress)', 'Maintenance')

doc.add_paragraph('Each service also shows notes from the IT team explaining any issues and the last update time.')

# ── 16. Contact & Support ───────────────────────────────────────────
doc.add_heading('16. Contact & Support', level=1)

doc.add_heading('Contact Form', level=2)
doc.add_paragraph(
    'Located at the Contact page, you can send a message directly to CAGD by filling out the contact form '
    'with your name, email, phone number (optional), subject, and message. An administrator will review '
    'and respond to your enquiry.'
)

doc.add_heading('Contact Information', level=2)
doc.add_paragraph('The contact page also displays:')
add_bullet('CAGD headquarters address')
add_bullet('Main phone number')
add_bullet('General email address')
add_bullet('Working hours')
add_bullet('Social media links (Facebook, YouTube, Twitter, Instagram, TikTok)')

# ── 17. FAQ ─────────────────────────────────────────────────────────
doc.add_heading('17. Frequently Asked Questions (FAQ)', level=1)
doc.add_paragraph(
    'The FAQ page provides answers to commonly asked questions about CAGD\'s services. Questions are '
    'organised by category and displayed in an accordion format — click on a question to expand and '
    'read the answer. Use the search bar to find specific topics.'
)

# ── 18. ChatBot ─────────────────────────────────────────────────────
doc.add_heading('18. ChatBot Assistant', level=1)
doc.add_paragraph(
    'A floating chat icon appears in the bottom-right corner of every page. Click it to open the '
    'CAGD ChatBot assistant.'
)
doc.add_paragraph('The ChatBot can help you with:')
add_bullet('Answers to frequently asked questions about CAGD services')
add_bullet('Information about e-services (GIFMIS, e-Payslip, TPRS)')
add_bullet('Contact details and office locations')
add_bullet('Navigation help — finding specific pages or information')
add_bullet('General enquiries about CAGD operations')

doc.add_paragraph(
    'The ChatBot first checks its built-in FAQ database for answers. If it cannot find a match, '
    'it uses AI to provide a helpful response. Type your question in the text box and press Enter or '
    'click the Send button.'
)

# ── 19. Accessibility ───────────────────────────────────────────────
doc.add_heading('19. Accessibility Features', level=1)

doc.add_heading('Dark Mode', level=2)
doc.add_paragraph(
    'Click the theme toggle icon (sun/moon) in the top navigation bar to switch between Light Mode, '
    'Dark Mode, and System (automatically matches your device\'s preference). Dark mode reduces '
    'eye strain in low-light environments.'
)

doc.add_heading('Language Support', level=2)
doc.add_paragraph(
    'The website supports English and Twi. Click the language switcher in the top navigation bar '
    'to change the interface language. News articles with Twi translations will display in the '
    'selected language.'
)

doc.add_heading('Mobile Responsive', level=2)
doc.add_paragraph(
    'The website is fully responsive and adapts to all screen sizes. On mobile devices, the navigation '
    'menu collapses into a hamburger menu (☰), and page layouts adjust for comfortable reading on smaller screens.'
)

doc.add_heading('Back to Top', level=2)
doc.add_paragraph(
    'When scrolling down on any page, a floating "Back to Top" button appears in the bottom-right corner. '
    'Click it to quickly scroll back to the top of the page.'
)

# ── 20. Keyboard Shortcuts ──────────────────────────────────────────
doc.add_heading('20. Keyboard Shortcuts', level=1)

table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.LEFT
hdr = table.rows[0]
hdr.cells[0].text = 'Shortcut'
hdr.cells[1].text = 'Action'
for cell in hdr.cells:
    for p in cell.paragraphs:
        p.runs[0].bold = True

shortcuts = [
    ('Ctrl + K (or Cmd + K on Mac)', 'Open Global Search'),
    ('Escape', 'Close search, modals, or dialogs'),
    ('Arrow Up / Down', 'Navigate search results'),
    ('Enter', 'Select highlighted search result'),
]
for i, (key, action) in enumerate(shortcuts):
    table.rows[i + 1].cells[0].text = key
    table.rows[i + 1].cells[1].text = action

doc.add_paragraph()
doc.add_paragraph()

# ── Footer ──────────────────────────────────────────────────────────
footer_text = doc.add_paragraph()
footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_text.add_run(
    '— End of User Guide —\n\n'
    'Controller & Accountant-General\'s Department\n'
    'Ministry of Finance, Accra, Ghana\n'
    'https://cagd.gov.gh'
)
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Page numbers ────────────────────────────────────────────────────
add_page_numbers(doc)

# ── Save ────────────────────────────────────────────────────────────
output_dir = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(output_dir, 'CAGD_Website_User_Guide.docx')
doc.save(output_path)
print(f'Saved: {output_path}')
