"""Generate CAGD Admin Panel Guide (.docx)"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page numbering helper ───────────────────────────────────────────
def add_page_numbers(doc, header_text='CAGD Admin Panel — Guide'):
    """Add page numbers to footer, starting from second page (cover page excluded)."""
    for section in doc.sections:
        section.different_first_page_header_footer = True
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.clear()

        run1 = p.add_run('Page ')
        run1.font.size = Pt(9)
        run1.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run2 = p.add_run()
        run2._r.append(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run3 = p.add_run()
        run3._r.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run4 = p.add_run()
        run4._r.append(fldChar2)

        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.clear()
        hr = hp.add_run(header_text)
        hr.font.size = Pt(8)
        hr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        hr.italic = True

        first_footer = section.first_page_footer
        first_header = section.first_page_header
        if first_footer.paragraphs:
            first_footer.paragraphs[0].clear()
        if first_header.paragraphs:
            first_header.paragraphs[0].clear()

# ── Styles ──────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x00, 0x52, 0x33)
    if level == 1:
        hs.font.size = Pt(22)
    elif level == 2:
        hs.font.size = Pt(16)
    else:
        hs.font.size = Pt(13)

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(f' — {text}')
    else:
        p.add_run(text)

def add_step(number, text):
    p = doc.add_paragraph()
    run = p.add_run(f'Step {number}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x52, 0x33)
    p.add_run(text)

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run(f'Note: {text}')
    run.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x55, 0x00)

def add_table(headers, rows):
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('CONTROLLER & ACCOUNTANT-GENERAL\'S DEPARTMENT')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x00, 0x52, 0x33)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Admin Panel — Complete Guide')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)

doc.add_paragraph()
url = doc.add_paragraph()
url.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = url.add_run('https://cagd.gov.gh/admin')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x00, 0x52, 0x33)

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('March 2026\nVersion 1.0\nICT Division — For Authorised Administrators Only')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc = [
    '',
    'PART A — EXECUTIVE SUMMARY',
    '  1. Overview',
    '  2. User Roles',
    '  3. Feature Quick Reference',
    '',
    'PART B — DETAILED GUIDE',
    '  4. Logging In',
    '  5. Dashboard',
    '  6. Hero Slides Manager',
    '  7. Homepage Manager',
    '  8. Pages Content Manager',
    '  9. FAQ Manager',
    '  10. News Manager',
    '  11. Reports Manager',
    '  12. Events Manager',
    '  13. Gallery Manager',
    '  14. Contact Messages',
    '  15. Leadership Manager',
    '  16. Divisions Manager',
    '  17. Projects Manager',
    '  18. Regional Offices Manager',
    '  19. Staff Directory Manager',
    '  20. Newsletter Subscriptions',
    '  21. Feedback',
    '  22. Service Status Manager',
    '  23. Forms Library Manager',
    '  24. User Management',
    '  25. Site Settings',
    '',
    'PART C — SECURITY & MAINTENANCE',
    '  26. Security Overview',
    '  27. Deployment & Updates',
]
for item in toc:
    if item == '':
        doc.add_paragraph()
    elif item.startswith('PART'):
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.bold = True
        run.font.size = Pt(12)
    else:
        p = doc.add_paragraph(item)
        p.runs[0].font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# PART A — EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════
pa = doc.add_paragraph()
pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = pa.add_run('PART A — EXECUTIVE SUMMARY')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x00, 0x52, 0x33)
doc.add_paragraph()

# ── 1. Overview ─────────────────────────────────────────────────────
doc.add_heading('1. Overview', level=1)
doc.add_paragraph(
    'The CAGD Admin Panel is a web-based content management system that allows authorised personnel to '
    'manage all content and settings on the cagd.gov.gh website. The admin panel is accessible at '
    'https://cagd.gov.gh/admin and requires authentication with a valid CAGD email address.'
)
doc.add_paragraph('The admin panel provides management capabilities for:')
add_bullet('Website content (news, events, reports, gallery, pages, FAQs)')
add_bullet('Homepage configuration (hero slides, statistics, e-services, leadership quote)')
add_bullet('Data management (staff directory, regional offices, forms library)')
add_bullet('Communication (contact messages, newsletter subscriptions, feedback)')
add_bullet('System configuration (announcement banner, service status, site settings)')
add_bullet('User management (admin and editor accounts)')

# ── 2. User Roles ───────────────────────────────────────────────────
doc.add_heading('2. User Roles', level=1)
doc.add_paragraph('The system has two user roles with different access levels:')

doc.add_heading('Administrator (admin)', level=2)
doc.add_paragraph('Full access to all features. Administrators can:')
add_bullet('Manage all content (news, events, reports, gallery, etc.)')
add_bullet('Configure site settings and announcement banners')
add_bullet('Manage user accounts and assign roles')
add_bullet('Access all admin pages')
add_bullet('Upload files and images')

doc.add_heading('Editor (editor)', level=2)
doc.add_paragraph('Limited access for content management. Editors can access:')
add_bullet('Dashboard (view only)')
add_bullet('News Manager (create, edit, delete articles)')
add_bullet('Reports Manager (create, edit, delete reports)')
add_bullet('Events Manager (create, edit, delete events and view registrations)')
add_bullet('Contact Messages (view and reply)')

doc.add_paragraph()
add_note('Editors cannot access User Management, Site Settings, Service Status, or any other admin pages not listed above.')

doc.add_heading('Access Summary Table', level=2)
add_table(
    ['Feature', 'Admin', 'Editor'],
    [
        ['Dashboard', 'Full', 'View Only'],
        ['Hero Slides', 'Yes', 'No'],
        ['Homepage', 'Yes', 'No'],
        ['Pages Content', 'Yes', 'No'],
        ['FAQs', 'Yes', 'No'],
        ['News', 'Yes', 'Yes'],
        ['Reports', 'Yes', 'Yes'],
        ['Events', 'Yes', 'Yes'],
        ['Gallery', 'Yes', 'No'],
        ['Contact Messages', 'Yes', 'Yes'],
        ['Leadership', 'Yes', 'No'],
        ['Divisions', 'Yes', 'No'],
        ['Projects', 'Yes', 'No'],
        ['Regional Offices', 'Yes', 'No'],
        ['Staff Directory', 'Yes', 'No'],
        ['Subscriptions', 'Yes', 'No'],
        ['Feedback', 'Yes', 'No'],
        ['Service Status', 'Yes', 'No'],
        ['Forms Library', 'Yes', 'No'],
        ['User Management', 'Yes', 'No'],
        ['Site Settings', 'Yes', 'No'],
    ]
)

# ── 3. Feature Quick Reference ──────────────────────────────────────
doc.add_heading('3. Feature Quick Reference', level=1)
add_table(
    ['Feature', 'URL', 'What It Does'],
    [
        ['Dashboard', '/admin', 'Statistics overview with charts and counters'],
        ['Hero Slides', '/admin/hero-slides', 'Manage homepage slideshow images and text'],
        ['Homepage', '/admin/homepage', 'Edit homepage sections (stats, functions, e-services, quote)'],
        ['Pages Content', '/admin/pages-content', 'Edit content for About, Division, and other pages'],
        ['FAQs', '/admin/faqs', 'Manage Frequently Asked Questions'],
        ['News', '/admin/news', 'Create/edit news articles with Twi translation'],
        ['Reports', '/admin/reports', 'Upload and manage downloadable reports'],
        ['Events', '/admin/events', 'Create events, manage registrations, export CSV'],
        ['Gallery', '/admin/gallery', 'Manage photo albums and images'],
        ['Messages', '/admin/messages', 'View and reply to contact form submissions'],
        ['Leadership', '/admin/leadership', 'Manage leadership profiles and bios'],
        ['Divisions', '/admin/divisions', 'Edit division information'],
        ['Projects', '/admin/projects', 'Manage project pages (PFMRP, IPSAS)'],
        ['Regional Offices', '/admin/regional-offices', 'Manage 16 regional office details and directors'],
        ['Staff Directory', '/admin/staff', 'Manage staff list with CSV/Excel import'],
        ['Subscriptions', '/admin/subscriptions', 'View newsletter subscribers, export CSV'],
        ['Feedback', '/admin/feedback', 'View event feedback and ratings, export CSV'],
        ['Service Status', '/admin/service-status', 'Set system status (GIFMIS, e-Payslip, etc.)'],
        ['Forms Library', '/admin/forms', 'Manage downloadable forms by category'],
        ['User Management', '/admin/users', 'Add admin/editor users'],
        ['Site Settings', '/admin/settings', 'Configure announcement banner and site settings'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# PART B — DETAILED GUIDE
# ══════════════════════════════════════════════════════════════════════
pa = doc.add_paragraph()
pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = pa.add_run('PART B — DETAILED GUIDE')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x00, 0x52, 0x33)
doc.add_paragraph()

# ── 4. Logging In ───────────────────────────────────────────────────
doc.add_heading('4. Logging In', level=1)
add_step(1, 'Open your browser and go to https://cagd.gov.gh/admin/login')
add_step(2, 'Enter your CAGD email address (must be @cagd.gov.gh)')
add_step(3, 'Enter your password')
add_step(4, 'Click "Sign In"')

doc.add_paragraph()
doc.add_heading('Forgot Password', level=2)
add_step(1, 'On the login page, click "Forgot Password?"')
add_step(2, 'Enter your email address')
add_step(3, 'Check your email for a password reset link')
add_step(4, 'Click the link and set a new password')

doc.add_paragraph()
add_note('Only accounts created by an administrator can log in. You cannot self-register.')

# ── 5. Dashboard ────────────────────────────────────────────────────
doc.add_heading('5. Dashboard', level=1)
doc.add_paragraph(
    'The Dashboard is the first page you see after logging in. It provides an overview of the website\'s content and activity.'
)

doc.add_heading('Statistics Cards', level=2)
doc.add_paragraph('The top section shows key metrics:')
add_bullet('Total News Posts')
add_bullet('Total Reports')
add_bullet('Total Events')
add_bullet('Gallery Albums')
add_bullet('Leadership Profiles')
add_bullet('Total Document Downloads')

doc.add_heading('Analytics Charts', level=2)
doc.add_paragraph(
    'Below the statistics, interactive charts show event registration data. Select a specific event '
    'from the dropdown to see:'
)
add_bullet('Registrations by Gender (pie chart)')
add_bullet('Registrations by Participant Type (pie chart)')
add_bullet('Registrations by Region (pie chart)')
add_bullet('Feedback Ratings Distribution (pie chart)')
add_bullet('Registration Trend Over Time (line chart)')

# ── 6. Hero Slides ──────────────────────────────────────────────────
doc.add_heading('6. Hero Slides Manager', level=1)
doc.add_paragraph(
    'Manages the large slideshow on the homepage. The website has 6 default hero images. '
    'You can create custom slides with your own titles, subtitles, and call-to-action buttons.'
)

doc.add_heading('Adding a New Slide', level=2)
add_step(1, 'Click "Add Slide"')
add_step(2, 'Select a hero image from the available options (hero-1 through hero-6)')
add_step(3, 'Enter a Title (displayed prominently on the slide)')
add_step(4, 'Enter a Subtitle (smaller text below the title)')
add_step(5, 'Optionally add a CTA Label (button text) and CTA Link (URL the button opens)')
add_step(6, 'Click "Save"')

doc.add_heading('Editing or Deleting Slides', level=2)
add_bullet('Click the pencil icon to edit an existing slide')
add_bullet('Click the trash icon to delete a slide')
add_bullet('Slides display in the order they are listed')

add_note('Hero images must be local files stored on the server at /images/hero/. Remote URLs are not supported for hero images.')

# ── 7. Homepage Manager ─────────────────────────────────────────────
doc.add_heading('7. Homepage Manager', level=1)
doc.add_paragraph('Allows you to customise the four main sections of the homepage:')

doc.add_heading('Statistics Counters', level=2)
doc.add_paragraph(
    'Edit the animated counter values shown on the homepage. Each counter has a label, value, and icon.'
)

doc.add_heading('Core Functions', level=2)
doc.add_paragraph('Edit the core functions cards with their titles, descriptions, and icons.')

doc.add_heading('e-Services', level=2)
doc.add_paragraph('Configure the e-Services section with service names, descriptions, URLs, and colours.')

doc.add_heading('Leadership Quote', level=2)
doc.add_paragraph('Set the quote text, author name, and title displayed on the homepage.')

add_note('Changes are saved to the database and appear on the website immediately.')

# ── 8. Pages Content ────────────────────────────────────────────────
doc.add_heading('8. Pages Content Manager', level=1)
doc.add_paragraph(
    'Manage the text content displayed on various public pages (About, Divisions, etc.) without '
    'needing to modify code. Select a page from the list, edit its content, and save.'
)

# ── 9. FAQ Manager ──────────────────────────────────────────────────
doc.add_heading('9. FAQ Manager', level=1)
doc.add_paragraph('Manage the questions and answers displayed on the public FAQ page and used by the ChatBot.')

doc.add_heading('Adding a FAQ', level=2)
add_step(1, 'Click "Add FAQ"')
add_step(2, 'Enter the Question')
add_step(3, 'Enter the Answer')
add_step(4, 'Select a Section/Category')
add_step(5, 'Click "Save"')

add_note('FAQs are also used by the ChatBot — when a user asks a question, the bot searches FAQs first before using AI.')

# ── 10. News Manager ────────────────────────────────────────────────
doc.add_heading('10. News Manager', level=1)
doc.add_paragraph('The most feature-rich admin section. Manage all news articles on the website.')

doc.add_heading('Creating a News Article', level=2)
add_step(1, 'Click "Add News"')
add_step(2, 'Enter the Title (English)')
add_step(3, 'The Slug (URL-friendly name) is auto-generated from the title')
add_step(4, 'Write an Excerpt (short summary shown on news listing pages)')
add_step(5, 'Write the full Content using the rich text editor (supports formatting, images, links, tables)')
add_step(6, 'Select a Category from the dropdown (General, Announcements, IPSAS, Payroll, GIFMIS, Training, Reforms, Press Release, Events, Digest, Treasury News)')
add_step(7, 'Upload a Featured Image (shown as the article thumbnail)')
add_step(8, 'Optionally attach a File (PDF download)')
add_step(9, 'Add Tags (comma-separated keywords)')
add_step(10, 'Set Status to "Published" (or keep as "Draft" to save without publishing)')
add_step(11, 'Click "Save"')

doc.add_heading('Twi Translation', level=2)
doc.add_paragraph(
    'Each article can be translated to Twi for bilingual visitors. After writing the English content:'
)
add_step(1, 'Click the "Translate to Twi" button')
add_step(2, 'The system automatically translates the title, excerpt, and content to Twi')
add_step(3, 'Review and edit the translations if needed')
add_step(4, 'Save the article')
add_note('You can also use "Bulk Translate All" to translate multiple articles at once.')

doc.add_heading('Managing Articles', level=2)
add_bullet('Use the Search bar to find articles by title')
add_bullet('Click column headers to sort by title, date, category, or status')
add_bullet('Click the pencil icon to edit an article')
add_bullet('Click the trash icon to delete an article')
add_bullet('Articles are paginated (15 per page)')

# ── 11. Reports Manager ─────────────────────────────────────────────
doc.add_heading('11. Reports Manager', level=1)
doc.add_paragraph('Upload and manage downloadable reports and publications.')

doc.add_heading('Adding a Report', level=2)
add_step(1, 'Click "Add Report"')
add_step(2, 'Enter the Report Title')
add_step(3, 'Select a Category')
add_step(4, 'Add a Description')
add_step(5, 'Upload the report file or paste a file URL')
add_step(6, 'Set the publication date')
add_step(7, 'Click "Save"')

doc.add_paragraph('Reports appear on the public Reports page and show download counts automatically.')

# ── 12. Events Manager ──────────────────────────────────────────────
doc.add_heading('12. Events Manager', level=1)
doc.add_paragraph('Create and manage events with registration capability.')

doc.add_heading('Creating an Event', level=2)
add_step(1, 'Click "Add Event"')
add_step(2, 'Enter the Event Title')
add_step(3, 'Write a Description (supports rich text)')
add_step(4, 'Set the Start Date and End Date')
add_step(5, 'Enter the Venue/Location')
add_step(6, 'Select a Category')
add_step(7, 'Upload a Featured Image')
add_step(8, 'Set Status (Published/Draft)')
add_step(9, 'Click "Save"')

doc.add_heading('Viewing Event Registrations', level=2)
doc.add_paragraph('For published events, you can view all registrations:')
add_step(1, 'Click the "View Registrations" button on the event row')
add_step(2, 'A dialog shows all registrants with their details')
add_step(3, 'You can update each registration\'s status (Pending, Confirmed, Cancelled)')
add_step(4, 'Click "Export CSV" to download all registrations as a spreadsheet')

doc.add_heading('Registration Data Includes', level=2)
add_bullet('Full name (first, middle, surname)')
add_bullet('Gender')
add_bullet('Participant type')
add_bullet('Region')
add_bullet('Department and organisation')
add_bullet('Phone and email')
add_bullet('Registration status')
add_bullet('Registration date/time')

# ── 13. Gallery Manager ─────────────────────────────────────────────
doc.add_heading('13. Gallery Manager', level=1)
doc.add_paragraph('Manage photo albums and images displayed in the public gallery.')

doc.add_heading('Creating an Album', level=2)
add_step(1, 'Click "Add Album"')
add_step(2, 'Enter the Album Title')
add_step(3, 'Upload a Cover Image')
add_step(4, 'Set the Album Date')
add_step(5, 'Optionally select a Parent Album (for nested albums/sub-albums)')
add_step(6, 'Click "Save"')

doc.add_heading('Adding Photos to an Album', level=2)
add_step(1, 'Click on an album to open it')
add_step(2, 'Click "Add Photos"')
add_step(3, 'Select one or more image files')
add_step(4, 'Photos upload with progress tracking')
add_step(5, 'Add captions to each photo (optional)')
add_step(6, 'Drag photos to reorder them')

# ── 14. Contact Messages ────────────────────────────────────────────
doc.add_heading('14. Contact Messages', level=1)
doc.add_paragraph(
    'View and manage messages submitted through the public contact form.'
)

doc.add_heading('Managing Messages', level=2)
add_bullet('New messages appear with an "Unread" indicator')
add_bullet('Click a message to view its full content')
add_bullet('Click "Mark as Read" to acknowledge it')
add_bullet('Use the Reply function to send an email response')
add_bullet('Use Search to find messages by name, email, or subject')
add_bullet('Messages are paginated (15 per page)')

# ── 15. Leadership Manager ──────────────────────────────────────────
doc.add_heading('15. Leadership Manager', level=1)
doc.add_paragraph('Manage the leadership profiles displayed on the Management → Leadership page.')

doc.add_heading('Adding a Profile', level=2)
add_step(1, 'Click "Add Profile"')
add_step(2, 'Enter Full Name')
add_step(3, 'Enter Title/Position')
add_step(4, 'Write a Biography')
add_step(5, 'Upload or paste a Photo URL')
add_step(6, 'Set the Display Order (lower numbers appear first)')
add_step(7, 'Select Profile Type (Leadership or Regional)')
add_step(8, 'Click "Save"')

# ── 16-17. Divisions & Projects ─────────────────────────────────────
doc.add_heading('16. Divisions Manager', level=1)
doc.add_paragraph(
    'Manage the content for CAGD\'s six division pages (Finance & Administration, Treasury, FMS, ICT, Payroll, Audit). '
    'Edit descriptions, staff information, and division-specific content.'
)

doc.add_heading('17. Projects Manager', level=1)
doc.add_paragraph(
    'Manage project pages such as PFMRP and IPSAS. Edit project descriptions, objectives, timelines, and related documents.'
)

# ── 18. Regional Offices ────────────────────────────────────────────
doc.add_heading('18. Regional Offices Manager', level=1)
doc.add_paragraph('Manage the 16 CAGD regional office details displayed on the Regional Offices and Regional Directors pages.')

doc.add_heading('Editing a Regional Office', level=2)
add_step(1, 'Click the pencil icon next to the region')
add_step(2, 'Update the Region Name, Address, Phone, and Email')
add_step(3, 'Enter or update the Director Name')
add_step(4, 'Upload a Director Photo (click "Upload Photo" and select an image file)')
add_step(5, 'Alternatively, paste a photo URL in the text field')
add_step(6, 'Click "Save"')

doc.add_heading('Director Photo Upload', level=2)
doc.add_paragraph(
    'Photos are uploaded directly to the server (not to external storage). Accepted formats: JPG, PNG, WebP, GIF. '
    'Maximum file size: 5 MB. Photos appear on the Regional Directors page and the Regional Offices Map.'
)

# ── 19. Staff Directory Manager ─────────────────────────────────────
doc.add_heading('19. Staff Directory Manager', level=1)
doc.add_paragraph(
    'Manage the staff directory that is searchable by the public. You can add staff one by one or import in bulk from a spreadsheet.'
)

doc.add_heading('Adding Staff Manually', level=2)
add_step(1, 'Click "Add Staff"')
add_step(2, 'Enter Full Name (required)')
add_step(3, 'Enter Job Title')
add_step(4, 'Enter Division and Department')
add_step(5, 'Enter Phone and Email')
add_step(6, 'Upload a Photo (optional)')
add_step(7, 'Set Order Position (controls display order; lower numbers appear first)')
add_step(8, 'Check "Active" to make the entry visible on the public directory')
add_step(9, 'Click "Add Staff Member"')

doc.add_heading('Importing Staff from CSV/Excel', level=2)
doc.add_paragraph(
    'You can bulk-import staff from a CSV or Excel (.xlsx) file. This is useful for importing from Active Directory or HR systems.'
)
add_step(1, 'Click the "Import" button')
add_step(2, 'Optionally download the Template (.xlsx) to see the expected column format')
add_step(3, 'Prepare your file with these columns: Name, Title, Division, Department, Phone, Email')
add_step(4, 'Click "Click to choose a CSV or Excel file" and select your file')
add_step(5, 'A preview table shows the first 5 rows — verify the data looks correct')
add_step(6, 'Click "Import X Staff Members" to import all records')

doc.add_heading('Accepted Column Names', level=3)
add_table(
    ['Your Column Name', 'Maps To'],
    [
        ['Name / Full Name / Fullname / Staff Name', 'Name'],
        ['Title / Job Title / Jobtitle / Position', 'Title'],
        ['Division', 'Division'],
        ['Department / Dept', 'Department'],
        ['Phone / Telephone / Mobile / Phone Number', 'Phone'],
        ['Email / Email Address', 'Email'],
    ]
)
add_note('Column matching is case-insensitive. Unrecognised columns are ignored.')

# ── 20. Subscriptions ───────────────────────────────────────────────
doc.add_heading('20. Newsletter Subscriptions', level=1)
doc.add_paragraph(
    'View all newsletter subscribers who signed up via the website\'s newsletter form (in the footer and news pages).'
)
add_bullet('View subscriber names and email addresses')
add_bullet('See subscription date')
add_bullet('Export all subscribers as CSV for use in email marketing tools')

# ── 21. Feedback ────────────────────────────────────────────────────
doc.add_heading('21. Feedback', level=1)
doc.add_paragraph('View feedback submitted by event attendees.')
add_bullet('See star ratings (1–5) and written feedback comments')
add_bullet('Filter by event')
add_bullet('View average ratings')
add_bullet('Export all feedback as CSV')

# ── 22. Service Status ──────────────────────────────────────────────
doc.add_heading('22. Service Status Manager', level=1)
doc.add_paragraph(
    'Control the public System Status page (/status) that shows the operational status of key government systems.'
)

doc.add_heading('Updating a Service Status', level=2)
add_step(1, 'Go to Admin → Service Status')
add_step(2, 'For each service (GIFMIS, e-Payslip, TPRS, CAGD Website), select a status:')
add_bullet('Operational — System is working normally (green)')
add_bullet('Degraded Performance — System is slow or partially available (amber)')
add_bullet('Service Outage — System is down (red)')
add_bullet('Under Maintenance — Planned maintenance (blue)')
add_step(3, 'Add a Note explaining the situation (optional but recommended during outages)')
add_step(4, 'Click "Save"')
add_note('Changes appear immediately on the public /status page.')

# ── 23. Forms Library ───────────────────────────────────────────────
doc.add_heading('23. Forms Library Manager', level=1)
doc.add_paragraph('Manage the downloadable forms available on the public Forms Library page.')

doc.add_heading('Adding a Form', level=2)
add_step(1, 'Click "Add Form"')
add_step(2, 'Enter the Form Name (e.g., "Pension Application Form")')
add_step(3, 'Enter a Description')
add_step(4, 'Select a Category (Payroll, Pension, GIFMIS, HR, Finance, General)')
add_step(5, 'Enter a Form Code (e.g., "CAGD-PEN-001")')
add_step(6, 'Paste the File URL (link to the downloadable file)')
add_step(7, 'Enter the File Size (e.g., "2.5 MB")')
add_step(8, 'Check "Active" to make it visible on the public page')
add_step(9, 'Click "Save"')

# ── 24. User Management ─────────────────────────────────────────────
doc.add_heading('24. User Management', level=1)
doc.add_paragraph(
    'Manage admin and editor accounts. Only users with the Admin role can access this page.'
)

doc.add_heading('Adding a New User', level=2)
add_step(1, 'Go to Admin → User Management')
add_step(2, 'Enter the user\'s CAGD email address (must end with @cagd.gov.gh)')
add_step(3, 'Enter a Password that meets the security requirements:')
add_bullet('Minimum 8 characters')
add_bullet('At least 1 uppercase letter (A-Z)')
add_bullet('At least 1 lowercase letter (a-z)')
add_bullet('At least 1 number (0-9)')
add_bullet('At least 1 special character (!@#$%^&*)')
add_step(4, 'Select a Role: Admin (full access) or Editor (limited access)')
add_step(5, 'Click "Add User"')

doc.add_heading('Editing User Roles', level=2)
doc.add_paragraph(
    'In the user list, you can change any user\'s role by selecting a new role from the dropdown in the '
    'Role column. The change takes effect immediately.'
)
add_note('You cannot change your own role (to prevent accidental self-demotion).')

# ── 25. Site Settings ───────────────────────────────────────────────
doc.add_heading('25. Site Settings', level=1)
doc.add_paragraph('Configure site-wide settings. Only accessible to Administrators.')

doc.add_heading('Announcement Banner', level=2)
doc.add_paragraph(
    'The announcement banner is a coloured strip that appears at the top of every public page. '
    'Use it for important notices, maintenance alerts, or deadline reminders.'
)
add_step(1, 'Go to Admin → Site Settings')
add_step(2, 'Find the Announcement Banner section')
add_step(3, 'Toggle "Active" to ON')
add_step(4, 'Enter the Banner Text')
add_step(5, 'Select a Type:')
add_bullet('Info (blue) — General information')
add_bullet('Warning (amber) — Important notices')
add_bullet('Error (red) — Urgent alerts')
add_bullet('Success (green) — Positive announcements')
add_step(6, 'Optionally add a Link URL and Link Label (creates a clickable button)')
add_step(7, 'Click "Save"')
add_note('To deactivate the banner, simply toggle "Active" to OFF and save.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# PART C — SECURITY & MAINTENANCE
# ══════════════════════════════════════════════════════════════════════
pa = doc.add_paragraph()
pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = pa.add_run('PART C — SECURITY & MAINTENANCE')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x00, 0x52, 0x33)
doc.add_paragraph()

doc.add_heading('26. Security Overview', level=1)
doc.add_paragraph('The CAGD website implements multiple layers of security:')

doc.add_heading('Authentication', level=2)
add_bullet('Only @cagd.gov.gh email addresses can be used for admin accounts')
add_bullet('Strong password requirements enforced (8+ characters, mixed case, numbers, special characters)')
add_bullet('JWT (JSON Web Token) based session management with automatic expiration')
add_bullet('Password reset via email verification')

doc.add_heading('Content Security', level=2)
add_bullet('Content-Security-Policy (CSP) header prevents cross-site scripting (XSS) attacks')
add_bullet('All user-generated HTML content is sanitised with DOMPurify before rendering')
add_bullet('Strict-Transport-Security (HSTS) enforces HTTPS connections')
add_bullet('X-Frame-Options prevents clickjacking')
add_bullet('X-Content-Type-Options prevents MIME sniffing')

doc.add_heading('File Upload Security', level=2)
add_bullet('Uploads require authenticated Supabase JWT token (not accessible to anonymous users)')
add_bullet('File type validation: server-side MIME detection, magic bytes verification, extension whitelist')
add_bullet('Only image files allowed (JPG, PNG, WebP, GIF) — maximum 5 MB')
add_bullet('Upload folders restricted to approved directories only')
add_bullet('Rate limiting: maximum 10 uploads per minute per IP address')
add_bullet('PHP tag detection prevents executable file uploads')

doc.add_heading('Access Control', level=2)
add_bullet('Role-based access control (Admin vs Editor) enforced in the UI')
add_bullet('Supabase Row Level Security (RLS) enforces access rules at the database level')
add_bullet('Sensitive admin pages (User Management, Site Settings, Service Status) restricted to Admin role only')
add_bullet('Directory listing disabled (Options -Indexes)')
add_bullet('Sensitive files (.env, .sql, .log, .bak) blocked from web access')

doc.add_heading('27. Deployment & Updates', level=1)
doc.add_paragraph('Website updates are deployed by the ICT Division. The deployment process:')
add_step(1, 'Code changes are made and tested locally')
add_step(2, 'The project is built (npm run build)')
add_step(3, 'Asset files (JS, CSS) are uploaded first')
add_step(4, 'Then index.html is uploaded last (prevents broken state during deployment)')
add_step(5, 'Deployment is verified by checking that the live site loads correctly')

doc.add_paragraph()
doc.add_heading('Technical Stack', level=2)
add_table(
    ['Component', 'Technology'],
    [
        ['Frontend Framework', 'React 18 + TypeScript'],
        ['Build Tool', 'Vite'],
        ['Styling', 'Tailwind CSS + shadcn/ui'],
        ['Database', 'Supabase (PostgreSQL) — self-hosted'],
        ['Hosting', 'cPanel shared hosting (LiteSpeed)'],
        ['Server', 'premium328.web-hosting.com'],
        ['Domain', 'cagd.gov.gh (DNS managed by NITA Ghana)'],
        ['ChatBot', 'Cloudflare Worker AI backend'],
        ['Email', 'Supabase Edge Functions + Resend'],
        ['Maps', 'Leaflet (OpenStreetMap)'],
        ['Charts', 'Recharts'],
        ['PWA', 'Service Worker with offline caching'],
    ]
)

# ── Footer ──────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
footer_text = doc.add_paragraph()
footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_text.add_run(
    '— End of Admin Guide —\n\n'
    'Controller & Accountant-General\'s Department\n'
    'Ministry of Finance, Accra, Ghana\n'
    'https://cagd.gov.gh\n\n'
    'CONFIDENTIAL — For Authorised Administrators Only'
)
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Page numbers ────────────────────────────────────────────────────
add_page_numbers(doc, 'CAGD Admin Panel — Guide  |  CONFIDENTIAL')

# ── Save ────────────────────────────────────────────────────────────
output_dir = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(output_dir, 'CAGD_Admin_Panel_Guide.docx')
doc.save(output_path)
print(f'Saved: {output_path}')
